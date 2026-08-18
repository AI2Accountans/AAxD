import docx
from docx import Document

def main():
    doc_path = 'Momento_0_Executive_Brief_EN.docx'
    doc = Document(doc_path)

    # Find the paragraph we want to modify
    target_text = "Fusing the operational expertise of Richard Gasca's 'First Mile' with Charles Hoffman's regulatory 'Last Mile' establishes a massive competitive advantage:"
    target_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if target_text.lower() in p.text.lower():
            target_idx = idx
            break

    print(f'Found target at paragraph index {target_idx}: {repr(doc.paragraphs[target_idx].text) if target_idx != -1 else "None"}')

    if target_idx != -1:
        # We will modify the target paragraph's text to the new introduction
        p_intro = doc.paragraphs[target_idx]
        p_intro.text = ''
        p_intro.paragraph_format.space_before = docx.shared.Pt(12)
        p_intro.paragraph_format.space_after = docx.shared.Pt(6)
        run = p_intro.add_run('To establish a truly transformative paradigm, this architecture consciously pulls together four key sets of complementary skills, experience, and domain expertise:')
        
        # We want to insert the four pillars and the transition paragraph.
        # We will insert them before the next paragraph in the list, 
        # which is the original Paragraph 77 (now index target_idx + 1).
        next_p = doc.paragraphs[target_idx + 1]
        
        pillars = [
            ('Philippe & the DFRNT Team (Technical Enablement & Graph Modeling): ', 'Serving as the core technical enabler, leveraging extensive experience with TerminusDB, enterprise semantic web architectures, and advanced data graph modeling.'),
            ('Jonathan Schmidt (Industrial Engineering & Lean Six Sigma): ', 'Optimizing accounting flows with his training as an Industrial Engineer, integrating Lean Six Sigma techniques, principles, and philosophies to design waste-free, high-efficiency ledger pipelines with rigorous internal control loops.'),
            ('Richard Gasca (Enterprise Systems & Operational Thought Leadership): ', "Bringing visionary conceptual design and deep operational enterprise expertise, architecting the 'First Mile' transactional REA graph model, and bridging corporate business realities with advanced database mechanics."),
            ('Charles Hoffman (Accounting Information Systems Generalist): ', 'An AIS pioneer who bridges the team by seeing the big picture from start (transaction entry) to finish (financial analysis of regulatory reports), leveraging deep research into the W3C Semantic Web Stack (RDF, OWL, SHACL, SKOS) to natively apply XBRL within modern reporting, auditing, ledger architectures, and cryptography (including digital distributed ledgers).')
        ]
        
        # Let's create the paragraphs.
        for bold_part, regular_part in pillars:
            new_p = next_p.insert_paragraph_before('', style='List Bullet')
            new_p.paragraph_format.space_after = docx.shared.Pt(3)
            r_bold = new_p.add_run(bold_part)
            r_bold.bold = True
            new_p.add_run(regular_part)
            
        # Now let's insert the transition paragraph before the benefits
        trans_p = next_p.insert_paragraph_before('Fusing these four pillars establishes a massive competitive advantage:', style='Normal')
        trans_p.paragraph_format.space_before = docx.shared.Pt(12)
        trans_p.paragraph_format.space_after = docx.shared.Pt(6)
        
        doc.save(doc_path)
        print('Saved successfully.')
    else:
        print('Target paragraph not found!')

if __name__ == '__main__':
    main()
