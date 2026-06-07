import os
import docx
from docx import Document

def main():
    doc_path = r'C:\Users\IPHIX\Documents\Projects\DFRNT\Documentacion\Memoria\Momento_0_Executive_Brief_EN.docx'
    
    if not os.path.exists(doc_path):
        print(f"Error: Word file not found at {doc_path}")
        return
        
    doc = Document(doc_path)

    # Let's search for the paragraph describing Richard's role
    target_start = "Richard Gasca (Enterprise Systems & Operational Thought Leadership):"
    target_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(target_start):
            target_idx = idx
            break

    print(f'Found Richard description at index {target_idx}: {repr(doc.paragraphs[target_idx].text) if target_idx != -1 else "None"}')

    if target_idx != -1:
        p = doc.paragraphs[target_idx]
        p.text = ''
        p.paragraph_format.space_after = docx.shared.Pt(3)
        r_bold = p.add_run('Richard Gasca (End-to-End Financial Information Supply Chain Integration): ')
        r_bold.bold = True
        p.add_run('Spanning the entire financial information supply chain—from raw transaction entry to final auditor visualization. Utilizing 10 years of experience with Altova MapForce and W3C standards, he designs and executes the automated pipelines that ingest operational data (Google Sheets/Excel), map it into XBRL Global Ledger (XBRL GL), remap it to supervisory taxonomies (XBRL FR), and render interactive HTML for continuous auditing, physically realizing the W3C/XBRL International integrated data flow.')
        
        doc.save(doc_path)
        print('Word file updated successfully.')
    else:
        # Check if the paragraph was already updated to the new title
        target_start_new = "Richard Gasca (End-to-End Financial"
        for idx, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith(target_start_new):
                target_idx = idx
                break
        if target_idx != -1:
            print("Richard description is already updated to the new End-to-End Financial Information Supply Chain Integration role.")
        else:
            print('Richard paragraph not found!')

if __name__ == '__main__':
    main()
